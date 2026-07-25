#!/usr/bin/env python3
"""Create a Word document from DEPLOYMENT.md content."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_document():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Configure heading styles
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Calibri Light'
        heading_style.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    
    # Title
    title = doc.add_heading('Deployment Guide', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Secure Multi-Branch Document Exchange System')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = doc.styles['Subtitle']
    
    doc.add_paragraph()  # Spacer
    
    # Prerequisites
    doc.add_heading('Prerequisites', level=1)
    prereqs = [
        'Node.js 18+ (20+ recommended)',
        'npm or bun',
        'Caddy (reverse proxy)',
        'PostgreSQL 14+ (required)',
        'Redis 6+ (optional, recommended for production)',
        'OpenSSL (for generating secrets)'
    ]
    for item in prereqs:
        doc.add_paragraph(item, style='List Bullet')
    
    # Quick Start
    doc.add_heading('Quick Start (Development)', level=1)
    
    doc.add_heading('1. Install PostgreSQL', level=2)
    
    doc.add_heading('Windows:', level=3)
    code = doc.add_paragraph()
    code.style = doc.styles['Normal']
    run = code.add_run('choco install postgresql\n# Or use Docker:\ndocker run -d --name postgres -e POSTGRES_PASSWORD=postgres -p 5432:5432 postgres:16')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading('macOS:', level=3)
    code = doc.add_paragraph()
    run = code.add_run('brew install postgresql@16\nbrew services start postgresql@16')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading('Linux:', level=3)
    code = doc.add_paragraph()
    run = code.add_run('sudo apt update\nsudo apt install postgresql postgresql-contrib\nsudo systemctl start postgresql')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading('2. Install Redis (Recommended)', level=2)
    doc.add_paragraph('Redis enables distributed rate limiting, session caching, and job deduplication.')
    
    doc.add_heading('Windows:', level=3)
    code = doc.add_paragraph()
    run = code.add_run('docker run -d --name redis -p 6379:6379 redis:7-alpine')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading('macOS:', level=3)
    code = doc.add_paragraph()
    run = code.add_run('brew install redis\nbrew services start redis')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading('Linux:', level=3)
    code = doc.add_paragraph()
    run = code.add_run('sudo apt update\nsudo apt install redis-server\nsudo systemctl start redis')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading('3. Create Database', level=2)
    code = doc.add_paragraph()
    run = code.add_run('psql -U postgres\nCREATE DATABASE secure_exchange;\\q')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading('4. Setup Application', level=2)
    code = doc.add_paragraph()
    run = code.add_run('npm install\ncp .env.example .env\nnpx prisma generate\nnpx prisma migrate dev --name init\ncd mini-services/exchange-hub\nnpm install\nnpm run dev\nnpm run dev')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    # Production Deployment
    doc.add_heading('Production Deployment', level=1)
    
    doc.add_heading('1. Setup PostgreSQL', level=2)
    code = doc.add_paragraph()
    run = code.add_run('sudo -u postgres psql\nCREATE USER secure_exchange WITH PASSWORD \'your-secure-password\';\nCREATE DATABASE secure_exchange OWNER secure_exchange;\nGRANT ALL PRIVILEGES ON DATABASE secure_exchange TO secure_exchange;\\q')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading('2. Generate Security Secrets', level=2)
    code = doc.add_paragraph()
    run = code.add_run('node -e "console.log(require(\'crypto\').randomBytes(32).toString(\'hex\'))"')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading('3. Configure Environment', level=2)
    doc.add_paragraph('Create .env with production values:')
    code = doc.add_paragraph()
    run = code.add_run('NODE_ENV=production\nDATABASE_URL=postgresql://secure_exchange:your-secure-password@localhost:5432/secure_exchange?schema=public\nDATABASE_POOL_SIZE=20\nDATABASE_LOG_LEVEL=error\nSESSION_SECRET=<generated-session-secret>\nSEED_SECRET=<generated-seed-secret>\nHUB_SERVER_TOKEN=<generated-hub-token>\nLOG_LEVEL=info')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading('4. Build and Start', level=2)
    code = doc.add_paragraph()
    run = code.add_run('npm run build\nnpx prisma migrate deploy\nnpm run start')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading('5. Configure Storage Backend', level=2)
    
    doc.add_heading('Local Storage (Development):', level=3)
    code = doc.add_paragraph()
    run = code.add_run('STORAGE_BACKEND=local\nSTORAGE_LOCAL_PATH=db/vault')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading('S3-Compatible Storage (Production):', level=3)
    code = doc.add_paragraph()
    run = code.add_run('STORAGE_BACKEND=s3\nS3_BUCKET=your-secure-exchange-vault\nS3_REGION=us-east-1\nS3_ACCESS_KEY_ID=your-access-key\nS3_SECRET_ACCESS_KEY=your-secret-key')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading('6. PostgreSQL Backup Strategy', level=2)
    code = doc.add_paragraph()
    run = code.add_run('# Automated daily backup (add to crontab)\n0 2 * * * pg_dump -U secure_exchange secure_exchange | gzip > /backups/secure_exchange_$(date +\\%Y\\%m\\%d).sql.gz\n\n# Retain backups for 7 days\n0 3 * * * find /backups -name "secure_exchange_*.sql.gz" -mtime +7 -delete')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading('7. S3 Backup Strategy', level=2)
    code = doc.add_paragraph()
    run = code.add_run('aws s3 sync s3://your-vault-bucket s3://your-backup-bucket --sse AES256\naws s3api put-bucket-versioning --bucket your-vault-bucket --versioning-configuration Status=Enabled')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    # Initial Setup
    doc.add_heading('Initial Setup', level=1)
    
    doc.add_heading('First Login', level=2)
    steps = [
        'Access the application at https://your-domain.com',
        'Use the OWNER credentials from the seed output',
        'Change the OWNER password immediately',
        'Enable 2FA for the OWNER account',
        'Create additional users with strong passwords'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('Seed Credentials', level=2)
    doc.add_paragraph('After seeding, credentials are returned in the API response. Store them securely:')
    roles = [
        'OWNER: Full system control',
        'SECURITY_ADMIN: User/branch/key management',
        'BRANCH_ADMIN: Branch-level management',
        'USER: Document send/receive',
        'READONLY: View-only access'
    ]
    for role in roles:
        doc.add_paragraph(role, style='List Bullet')
    
    # Backup and Recovery
    doc.add_heading('Backup and Recovery', level=1)
    
    doc.add_heading('Creating Backups', level=2)
    code = doc.add_paragraph()
    run = code.add_run('curl -X POST http://localhost:3000/api/backup \\\n  -H "Content-Type: application/json" \\\n  -H "Cookie: secure-exchange-session=<owner-token>" \\\n  -d \'{"action": "create"}\'')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading('Restoring from Backup', level=2)
    doc.add_paragraph('Backups are stored in db/backups/. To restore:')
    steps = [
        'Stop all services',
        'Copy the backup database over the current one',
        'Copy vault files',
        'Restart services'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    # Monitoring
    doc.add_heading('Monitoring', level=1)
    
    doc.add_heading('Health Check', level=2)
    code = doc.add_paragraph()
    run = code.add_run('curl http://localhost:3000/api/health')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_paragraph('Returns:')
    items = [
        'Database connectivity status',
        'System state (active/lockdown)',
        'Memory usage',
        'Uptime'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Logs', level=2)
    doc.add_paragraph('Logs are structured JSON. Use a log aggregator (ELK, Datadog, etc.) for production.')
    doc.add_paragraph('Environment variables:')
    items = [
        'LOG_LEVEL=info (default)',
        'LOG_LEVEL=debug (verbose)'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    # Security Checklist
    doc.add_heading('Security Checklist', level=1)
    checklist = [
        'All secrets generated and stored securely',
        '.env not committed to version control',
        'HTTPS enabled via Caddy',
        'Database credentials not in source code',
        'Owner password changed from seed default',
        '2FA enabled for all admin accounts',
        'Regular backups configured',
        'Monitoring and alerting configured',
        'Audit logs reviewed regularly',
        'Key rotation schedule established'
    ]
    for item in checklist:
        doc.add_paragraph(f'☐ {item}')
    
    # Troubleshooting
    doc.add_heading('Troubleshooting', level=1)
    
    doc.add_heading('"Seeding is disabled in production"', level=2)
    doc.add_paragraph('Set SEED_SECRET environment variable and authenticate as OWNER.')
    
    doc.add_heading('"Invalid seed secret"', level=2)
    doc.add_paragraph('Ensure the x-seed-secret header matches the SEED_SECRET env var.')
    
    doc.add_heading('Hub connection failed', level=2)
    doc.add_paragraph('Check that:')
    items = [
        'Exchange Hub is running on port 3003',
        'HUB_SERVER_TOKEN matches between Next.js and Hub',
        'Caddy is forwarding correctly'
    ]
    for i, item in enumerate(items, 1):
        doc.add_paragraph(f'{i}. {item}')
    
    doc.add_heading('Database errors', level=2)
    doc.add_paragraph('Check:')
    items = [
        'DATABASE_URL is correct',
        'Database server is running',
        'User has appropriate permissions'
    ]
    for i, item in enumerate(items, 1):
        doc.add_paragraph(f'{i}. {item}')
    
    # Scaling
    doc.add_heading('Scaling', level=1)
    
    doc.add_heading('Horizontal Scaling', level=2)
    doc.add_paragraph('For multiple Next.js instances:')
    items = [
        'Use PostgreSQL instead of SQLite',
        'Use Redis for session storage',
        'Use Redis for rate limiting',
        'Share the Exchange Hub connection'
    ]
    for i, item in enumerate(items, 1):
        doc.add_paragraph(f'{i}. {item}')
    
    doc.add_heading('Load Balancing', level=2)
    doc.add_paragraph('Configure Caddy or nginx to load balance across multiple Next.js instances.')
    
    # Support
    doc.add_heading('Support', level=1)
    doc.add_paragraph('For issues, check:')
    items = [
        'Application logs',
        'Database connectivity',
        'Hub connectivity',
        'Caddy configuration'
    ]
    for i, item in enumerate(items, 1):
        doc.add_paragraph(f'{i}. {item}')
    
    # Save document
    doc.save('DEPLOYMENT.docx')
    print('Created DEPLOYMENT.docx')

if __name__ == '__main__':
    create_document()
